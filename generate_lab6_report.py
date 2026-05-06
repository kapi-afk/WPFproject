# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT_PATH = Path(r"D:\WPF_project\Лабораторная работа №6_сервисный центр.docx")


def set_run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=14, indent=True):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    if indent and text:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_table(doc, rows_data, cols):
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"

    for row_index, row_data in enumerate(rows_data):
        row = table.add_row().cells
        for col_index, text in enumerate(row_data):
            paragraph = row[col_index].paragraphs[0]
            run = paragraph.add_run(text)
            set_run_font(run, size=12, bold=(row_index == 0))
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)

    add_paragraph(doc, "Белорусский государственный технологический университет", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph(doc, "Факультет информационных технологий", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph(doc, "Специальность программная инженерия", WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    for _ in range(5):
        add_paragraph(doc, "", indent=False)

    add_paragraph(doc, "Отчёт по лабораторной работе №6", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, indent=False)
    add_paragraph(doc, "По дисциплине «Разработка и анализ требований»", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph(doc, "На тему «Управление требованиями»", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph(doc, "для приложения сервисного центра по ремонту техники", WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    for _ in range(4):
        add_paragraph(doc, "", indent=False)

    add_paragraph(doc, "Выполнил:", indent=False)
    add_paragraph(doc, "Студент ________________________________", indent=False)
    add_paragraph(doc, "Преподаватель: ________________________", indent=False)

    for _ in range(5):
        add_paragraph(doc, "", indent=False)

    add_paragraph(doc, "2026, Минск", WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    doc.add_page_break()

    add_paragraph(doc, "Задание 1. Матрица трассировки требований.", bold=True, indent=False)
    add_paragraph(doc, "Тест-кейсы", bold=True, indent=False)

    test_cases = [
        "TC-1.1: Проверка успешного входа пользователя по логину и паролю.",
        "TC-1.2: Проверка регистрации нового клиента с уникальными логином и email.",
        "TC-1.3: Проверка отображения сообщения об ошибке при неверных учетных данных.",
        "TC-2.1: Проверка создания заявки на ремонт при заполнении обязательных полей.",
        "TC-2.2: Проверка запрета создания заявки для неавторизованного пользователя.",
        "TC-2.3: Проверка автоматического назначения мастера при наличии подходящего специалиста.",
        "TC-3.1: Проверка отображения списков пользователей, товаров и услуг в панели администратора.",
        "TC-3.2: Проверка изменения мастером статуса заказа.",
    ]
    for item in test_cases:
        add_paragraph(doc, item, indent=False)

    add_paragraph(doc, "Матрица трассировки требований", bold=True, indent=False)
    add_paragraph(doc, "Requirement number: Уникальный номер требования.", indent=False)
    add_paragraph(doc, "Module number: Модуль приложения.", indent=False)
    add_paragraph(doc, "High level requirement: Высокоуровневое требование.", indent=False)
    add_paragraph(doc, "Low level requirement: Детализированное требование.", indent=False)
    add_paragraph(doc, "Test case name: Связанный тест-кейс.", indent=False)

    traceability_rows = [
        ["Requirement number", "Module number", "High level requirement", "Low level requirement", "Test case name"],
        ["1", "1 Authentication", "1.1 Управление учетной записью клиента", "1.1.1 --> Вход пользователя по логину и паролю", "TC-1.1"],
        ["2", "", "", "1.1.2 --> Регистрация нового клиента", "TC-1.2"],
        ["3", "", "", "1.1.3 --> Отображение ошибки при неверных учетных данных", "TC-1.3"],
        ["4", "2 Repair order", "2.1 Оформление заявки на ремонт", "2.1.1 --> Создание заявки с обязательными полями", "TC-2.1"],
        ["5", "", "", "2.1.2 --> Запрет создания заявки без авторизации", "TC-2.2"],
        ["6", "", "", "2.1.3 --> Автоматическое назначение мастера", "TC-2.3"],
        ["7", "3 Administration", "3.1 Администрирование и обработка заказов", "3.1.1 --> Просмотр администратором списков пользователей, товаров и услуг", "TC-3.1"],
        ["8", "", "", "3.1.2 --> Изменение мастером статуса заказа", "TC-3.2"],
    ]
    add_table(doc, traceability_rows, cols=5)

    add_paragraph(doc, "", indent=False)
    add_paragraph(doc, "Задание 2. Анализ влияния изменений.", bold=True, indent=False)
    add_paragraph(doc, "Добавление незначительного изменения в требование", bold=True, indent=False)
    add_paragraph(doc, "Исходное требование (FR-2.1.1):", bold=True, indent=False)
    add_paragraph(
        doc,
        "Для создания заявки на ремонт пользователь должен указать тип устройства, бренд, модель, описание проблемы и контактный телефон.",
        indent=False,
    )
    add_paragraph(doc, "Измененное требование:", bold=True, indent=False)
    add_paragraph(
        doc,
        "Для создания заявки на ремонт пользователь должен указать тип устройства, бренд, модель, описание проблемы и контактный телефон в формате +375 XX XXX-XX-XX.",
        indent=False,
    )
    add_paragraph(doc, "Анализ влияния с помощью матрицы трассировки требований", bold=True, indent=False)

    impact_rows = [
        ["Объект проекта", "Действие по изменению"],
        ["Функциональное требование FR-2.1.1", "Обновить описание требования: добавить обязательную проверку формата телефона."],
        ["Тест-кейс TC-2.1", "Обновить проверку: заявка считается валидной только при корректном формате номера."],
        ["UI/UX дизайн: форма создания заявки", "Обновить макет поля телефона и добавить подсказку с требуемым форматом."],
    ]
    add_table(doc, impact_rows, cols=2)

    add_paragraph(doc, "", indent=False)
    add_paragraph(doc, "Задание 3. Работа с RACI-матрицей", bold=True, indent=False)
    add_paragraph(
        doc,
        "Ниже приведён пример RACI-матрицы для процесса реализации функции создания заявки на ремонт.",
        indent=False,
    )

    raci_rows = [
        ["Роль/задача", "Владелец продукта", "Разработчик", "Тестировщик", "Дизайнер"],
        ["Определение требований", "R", "C", "C", "C"],
        ["Проектирование UI", "A", "C", "-", "R"],
        ["Реализация формы заявки", "A", "R", "C", "-"],
        ["Настройка валидации полей", "C", "R", "A", "-"],
        ["Тестирование сценария", "I", "C", "R", "-"],
        ["Исправление дефектов", "-", "R", "A", "-"],
        ["Внедрение в приложение", "R", "A", "C", "-"],
    ]
    add_table(doc, raci_rows, cols=5)

    add_paragraph(doc, "", indent=False)
    add_paragraph(doc, "Что отражено в матрице:", bold=True, indent=False)
    add_paragraph(
        doc,
        "Матрица показывает распределение ответственности между участниками проекта при анализе, проектировании, реализации и тестировании функции создания заявки на ремонт.",
        indent=False,
    )
    add_paragraph(doc, "Проверка по горизонтали:", bold=True, indent=False)
    add_paragraph(
        doc,
        "В каждой строке есть исполнитель работы (R), а ключевые этапы имеют ответственного за результат (A), следовательно распределение ролей выполнено корректно.",
        indent=False,
    )
    add_paragraph(doc, "Проверка по вертикали:", bold=True, indent=False)
    add_paragraph(
        doc,
        "Пустых ролей нет: разработчик и тестировщик задействованы наиболее активно, дизайнер отвечает за интерфейс, а владелец продукта участвует в требованиях и внедрении.",
        indent=False,
    )

    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUT_PATH)
    print(OUT_PATH)
